import os
from flask import Flask, request, render_template,jsonify
import requests
from werkzeug.utils import secure_filename 
from xml.etree import ElementTree
from docx import Document
from docx.document import Document as _Document
import docx
from base64 import b64decode
import re
import json  
from datetime import date
import copy
import xml.etree.ElementTree as ET
from docx.enum.table import WD_ROW_HEIGHT
from Models.objwrapper import obj_wrap
from Models.fieldwrap import field_wrap_obj
from Models.grandwrap import grand_wrap_obj
from Models.parentwrap import parent_wrap_obj
from Models.childwrap import child_wrap_obj
from _io import BytesIO, StringIO
import base64
import io
import random
import string
import uuid 
import mysql.connector






UPLOAD_FOLDER = str(os.getcwd())+'/python-mailmerge/Document'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}
path = str(os.getcwd())+'/Document'
app = Flask(__name__)
mydb = mysql.connector.connect(host='localhost',user='root',password='Aspi@2018',auth_plugin='mysql_native_password')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER 
file_path = ''
document_data = ''
session_id = ''
record_id = ''

@app.route('/')
def index():
    return render_template('indexlocal.html')
    

@app.route('/handle_form', methods=['POST'])
def handle_form():
    
    # file = request.files['file']
    # filename = secure_filename(file.filename)
    # file_path = os.path.join(path, filename)
    # file.save(file_path)
    # doc = docx.Document(file_path)
    
     
    
    filename = 'quote.docx'
    file_path = os.path.join(path, filename)
    content = request.data
    # create_folder('./Document1/')
    folder_path = str(os.getcwd())
    id = uuid.uuid1()
    os.makedirs('./'+id.hex)
    # session_id = request.headers['Session-Id']
    session_id = '00DN0000000Xus2MAC!AR0AQB7ICQROcGsf6X_a1zamtLOJ1MLO9Ut1a1ndex3NfUbRQV.HePM9mTGyP.YNv_eNewYtp_rlIBBukTXOykCdhkCh57i1'
    instance_url = request.headers['baseUrl']
    record_id = request.headers['recordId']
    print("headers-->{}".format(request.data))
    bytes = b64decode(content)
    # f = open(file_path,'wb')
    source_stream = BytesIO(content)
    doc = Document(source_stream)
    source_stream.close()
    # f.write(bytes)
    # f.close
    # doc = docx.Document(documentx)
    full_text = []
    field_list = []
    child_obj_metadata = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    document_data = '\n'.join(full_text)
    field_list = re.findall("\\$\\{(.*?)\\}", document_data)
    child_obj_metadata = re.findall("\\$tbl\\{(.*?)\\}", document_data)
    head_child_obj = ''  
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                fields_in_cell = re.findall("\\$\\{(.*?)\\}",cell.text)
                child_obj_values = re.findall("\\$tbl\\{(.*?)\\}",cell.text)
                child_obj_str = []
                if len(child_obj_values) > 0:
                   child_obj_str =  re.findall("\\$tbl{START:(.*):", cell.text) 
                if len(child_obj_str) > 0:
                   head_child_obj = child_obj_str[0]
                    
                if len(fields_in_cell) > 0 and fields_in_cell[0].split('.')[1] == head_child_obj.strip():
                    for field in fields_in_cell :
                        child_obj_metadata.append(field)
                elif len(fields_in_cell) > 0 :
                    for field in fields_in_cell :
                        field_list.append(field)
    field_list = list(dict.fromkeys(field_list))
    
    #This block will do json formation of fields and parent objects which will be returned to salesforce for retrieving data             
    if len(field_list) > 0 :
        obj_wrapper = obj_wrap(field_list[0].split('.')[0],False,[],[],[])
        field_wrapper = []
        parent_wrapper = []
        parent_field_wrapper = []
        grand_parent_field_wrapper = []
        grand_wrapper = []
        for field in field_list:
            format_type = re.findall("(#[A-Z]*)",field)
            if len(format_type) > 0 :
                corrected_field = field.replace(format_type[0],'')
            else :
                corrected_field = field
            parent_obj = corrected_field.split('.')
            if len(parent_obj) == 2:
                field_wrap = field_wrap_obj(parent_obj[-1],False)
                if field_wrap not in field_wrapper:
                    field_wrapper.append(field_wrap)

            else:
                field_wrap = field_wrap_obj(parent_obj[1],False)
                if field_wrap.__dict__ not in field_wrapper:
                    field_wrapper.append(field_wrap.__dict__)
                filtered_obj = parent_obj[1:len(parent_obj)]

                if len(filtered_obj) == 3:
                    parent_field_wrap = field_wrap_obj(filtered_obj[1],False)
                    if parent_field_wrap not in parent_field_wrapper:
                        parent_field_wrapper.append(parent_field_wrap)
                    grand_parent_field_wrap = field_wrap_obj(filtered_obj[2],False)
                    if grand_parent_field_wrap not in grand_parent_field_wrapper:
                        grand_parent_field_wrapper.append(grand_parent_field_wrap)
                    grand_wrap = grand_wrap_obj(filtered_obj[1],False,grand_parent_field_wrapper)
                    if len(grand_wrapper) > 0:
                        check_grobj_list = list()
                        for obj in grand_wrapper:
                            check_grobj_list.append(obj.objName)

                        if grand_wrap.objName not in check_grobj_list:
                            grand_wrapper.append(grand_wrap)

                        elif {
                                'fieldName': filtered_obj[2]
                        } not in grand_wrapper[check_grobj_list.index(
                                grand_wrap.objName)].fieldWrapperList:
                            grand_wrapper[check_grobj_list.index(
                                grand_wrap.objName)].fieldWrapperList.append(
                                    field_wrap_obj(filtered_obj[2],False))
                    else:
                        grand_wrapper.append(grand_wrap)
                    parent_wrap = parent_wrap_obj(filtered_obj[0],False,parent_field_wrapper,[],[],grand_wrapper)

                    if len(parent_wrapper) > 0:
                        check_obj_list = list()
                        for obj in parent_wrapper:
                            check_obj_list.append(obj.objName)
                        if parent_wrap.objName not in check_obj_list:
                            parent_wrapper.append(parent_wrap)

                        else:

                            if field_wrap_obj(filtered_obj[1],False) not in parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName)].fieldWrapperList:

                                parent_wrapper[check_obj_list.index(
                                    parent_wrap.objName
                                )].fieldWrapperList.append(field_wrap_obj(filtered_obj[1],False))
                                if 'grandObjWrapperList' in parent_wrapper[
                                        check_obj_list.index(
                                            parent_wrap.objName)]:
                                    parent_wrapper[check_obj_list.index(
                                        parent_wrap.objName
                                    )].grandObjWrapperList.append(grand_wrap_obj(filtered_obj[1],False,[field_wrap_obj(filtered_obj[2],False)]))

                                else:
                                    parent_wrapper[check_obj_list.index(
                                        parent_wrap.objName
                                    )].grandObjWrapperList = [grand_wrap_obj(filtered_obj[1],False,[field_wrap_obj(filtered_obj[2],False)])]

                            else:

                                check_grandobj_list = list()
                                if 'grandObjWrapperList' in parent_wrapper[
                                        check_obj_list.index(
                                            parent_wrap.objName)]:

                                    for obj in parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                    )].grandObjWrapperList:
                                        check_grandobj_list.append(obj.objName)
                                else:
                                    check_grandobj_list = []
                                if grand_wrap.objName not in check_grandobj_list:
                                    grand_wrapper.append(grand_wrap)

                                else:

                                    if field_wrap_obj(filtered_obj[2],False) not in parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                    )].grandObjWrapperList[
                                            check_grandobj_list.index(
                                                grand_wrap.objName
                                            )].fieldWrapperList:

                                        parent_wrapper[check_obj_list.index(
                                            parent_wrap.objName
                                        )].grandObjWrapperList[
                                            check_grandobj_list.index(
                                                grand_wrap.objName
                                            )].fieldWrapperList.append(field_wrap_obj(filtered_obj[2],False))

                    else:
                        parent_wrapper.append(parent_wrap)

                    grand_wrap = {}
                    grand_parent_field_wrap = {}
                    parent_field_wrap = {}
                    parent_wrap = {}
                    parent_field_wrapper = []
                    grand_parent_field_wrapper = []
                    grand_wrapper = []

                elif len(filtered_obj) == 2:
                    parent_field_wrap = field_wrap_obj(filtered_obj[-1],False)
                    check_parent_obj_list = list()
                    for obj in parent_wrapper:
                        check_parent_obj_list.append(obj.objName)
                    if ({
                            'objName': filtered_obj[0]
                    } in check_parent_obj_list
                        ) and parent_field_wrap not in parent_wrapper[
                            check_parent_obj_list.index(
                                filtered_obj[0])]['fieldWrapperList']:
                        parent_field_wrapper.append(parent_field_wrap)

                    parent_wrap = parent_wrap_obj(filtered_obj[0],False,parent_field_wrapper,[],[],[])
                    if len(parent_wrapper) > 0:
                        check_obj_list = list()
                        for obj in parent_wrapper:
                            check_obj_list.append(obj.objName)
                        if parent_wrap.objName not in check_obj_list:
                            parent_wrap.fieldWrapperList = [field_wrap_obj(filtered_obj[-1],False)]
                            parent_wrapper.append(parent_wrap)
                            parent_field_wrapper = []

                        elif {
                                'fieldName': filtered_obj[-1],
                                'isExists': False
                        } not in parent_wrapper[check_obj_list.index(
                                parent_wrap.objName)].fieldWrapperList:
                            parent_wrapper[check_obj_list.index(
                                parent_wrap.objName
                            )].fieldWrapperList.append(field_wrap_obj(filtered_obj[-1],False))
                            parent_field_wrapper = []
                    else:
                        parent_wrap.fieldWrapperList.append(field_wrap_obj(filtered_obj[-1],False))
                        parent_wrapper.append(parent_wrap)
                        parent_field_wrapper = []
                parent_field_wrapper = []
        
        obj_wrapper.fieldWrapperList = field_wrapper
        obj_wrapper.parentObjWrapperList = parent_wrapper
        
        
       
        
        
    parent_wrapper = [] 
    parent_field_wrapper = []
    parent_field_wrap = {}
    parent_wrap = {}
    old_child_obj_meta = []
    
        
    
        
    
    #Method to return index of the obj 
    def check_obj_present(current_obj, child_wrapper):
        obj_list = list()
        for record in child_wrapper.parentObjWrapperList:
            obj_list.append(record.objName)
        if len(obj_list) > 0:
            try:
                return obj_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    
    #Method to return index of the parent obj
    def check_grand_obj_present(current_obj, child_wrapper, parent_index):
        obj_list = list()

        for record in child_wrapper.parentObjWrapperList[parent_index].grandObjWrapperList:
            obj_list.append(record.objName)
        if len(obj_list) > 0:
            try:
                return obj_list.index(current_obj)
            except ValueError:
                return -1
        else:
            return -1
    
    #Method to generate childWrapper
    def generate_child_obj(child_obj, child_wrapper):
        child1_field_wrap = field_wrap_obj(child_obj[2],False)
        if child1_field_wrap not in child_wrapper.fieldWrapperList:
            child_wrapper.fieldWrapperList.append(child1_field_wrap)
        if len(child_obj) == 4:
            index_value = check_obj_present(child_obj[2], child_wrapper)
            if index_value == -1:
                child_wrapper.parentObjWrapperList.append(parent_wrap_obj(child_obj[2],False,field_wrap_obj(child_obj[3],False),[],[]))
            else:
                if field_wrap_obj(child_obj[3],False) not in child_wrapper.parentObjWrapperList[check_obj_present(child_obj[2],child_wrapper)].fieldWrapperList:
                    child_wrapper.parentObjWrapperList[check_obj_present(child_obj[2],child_wrapper)].fieldWrapperList.append(field_wrap_obj(child_obj[3],False))
        if len(child_obj) == 5:
            index_value = check_obj_present(child_obj[2], child_wrapper)
            if index_value == -1:
                child_wrapper.parentObjWrapperList.append(parent_wrap_obj(child_obj[2],False,
                [field_wrap_obj(child_obj[3],False)],
                [grand_wrap_obj(child_obj[3],False,[field_wrap_obj(child_obj[4],False)])]
                ))
            else:
                if field_wrap_obj(child_obj[3],False) not in child_wrapper.parentObjWrapperList[index_value].fieldWrapperList:
                    child_wrapper.parentObjWrapperList[index_value].fieldWrapperList.append(field_wrap_obj(child_obj[3],False))
                grand_index_value = check_grand_obj_present(child_obj[3], child_wrapper, index_value)
                if grand_index_value == -1:
                    child_wrapper.parentObjWrapperList[index_value].grandObjWrapperList.append(grand_wrap_obj(child_obj[3],False,[field_wrap_obj(child_obj[4],False)]))
                else:
                    child_wrapper.parentObjWrapperList[index_value].grandObjWrapperList[grand_index_value].fieldWrapperList.append(field_wrap_obj(child_obj[4],False))
        return child_wrapper
    
    
            
    if len(child_obj_metadata) > 0:
        child_wrapper = child_wrap_obj('',False,[],[],'')
        child_obj_list = []
        child_wrapper_list = []
        for field in child_obj_metadata:
            child_obj = field.split('.')
            if child_obj not in old_child_obj_meta:
                if child_obj[1] not in child_obj_list:
                    head_obj = re.findall("\\$tbl{START:[A-Za-z]\\:(.*)",document_data)
                    child_obj_check = child_wrap_obj(child_obj[1],False,[],[],'')
                    child_wrapper = generate_child_obj(
                        child_obj,child_obj_check )
                    child_obj_list.append(child_obj[1])
                    child_wrapper_list.append(child_wrapper)
                else:
                    child_wrapper = generate_child_obj(
                        child_obj,
                        child_wrapper_list[child_obj_list.index(child_obj[1])])
                    child_wrapper_list[child_obj_list.index(child_obj[1])] = child_wrapper
                old_child_obj_meta.append(child_obj)
        obj_wrapper.childObjWrapperList = child_wrapper_list
        obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
        r = requests.post(instance_url+"/services/apexrest/data_retrieve/"+record_id,data=obj_wrapper,headers = {"Authorization":'OAuth '+session_id,'Content-Type': 'application/json','recordId':record_id})
        print("response-->{}".format(r))
        data_dict = json.loads(r.json())
        
        
        # r =  get_data_sf()
        # data_dict = json.loads(r.text)
        bind_values_doc(data_dict,doc)
        docx_stream = io.BytesIO()
        doc.save(docx_stream)
        docx_bytes = docx_stream.getvalue()
        encoded = base64.b64encode(docx_bytes)
        data = {
                "Name" : "demoAttachment.pdf",
                "Body": str(encoded)[2:-1],
                "parentId": record_id 
                }
        # salesforce_response = requests.post(instance_url+"/services/data/v47.0/sobjects/Attachment",data=data,headers = {"Authorization":'OAuth '+'00D0p0000000V5HEAU!AQgAQEWkzrkC5T1NKhmV2C43BfSYR.2NbZ2VxatnomppLnH9V6hnFl1SCwIun9Cm2FI9Xbdpt_Lp6ie.pUyXXuFKfFisMEiJ','Content-Type': 'application/json'})
        # print("salesforce_response-->{}".format(salesforce_response))
        return json.dumps(data)
    else :
        obj_wrapper = json.dumps(obj_wrapper, default=lambda o: o.__dict__)
        print("ObjMetaDataInfo-->{}".format(obj_wrapper))
        r = requests.post(instance_url+"/services/apexrest/data_retrieve/"+record_id,data=obj_wrapper,headers = {"Authorization":'OAuth '+session_id,'Content-Type': 'application/json','recordId':record_id})
        data_dict = json.loads(r.json())
        print("session_id-->{}".format(data_dict))
        
        bind_values_doc(data_dict,doc)
        docx_stream = io.BytesIO()
        doc.save(docx_stream)
        docx_bytes = docx_stream.getvalue()
        encoded = base64.b64encode(docx_bytes)
        data = {
                "Name" : "demoAttachment.pdf",
                "Body": str(encoded)[2:-1],
                "parentId": record_id 
                }
        # r =  get_data_sf()
        # data_dict = json.loads(r.text)
        # bind_values_doc(data_dict,file_path)
        return json.dumps(data)


#Create folder for document :
def create_folder(directory):

    try:
        folder_path = str(os.getcwd())
        os.chdir(folder_path)
        print("folder_path-->{}".format(folder_path))
        os.makedirs('Docu')
        print("Folder Created->")
        if not os.path.exists(directory):

            os.makedirs(directory)
    except OSError:
        print("Error: Creating directory. {}".format(directory))
@app.route('/get_document',method='POST')


def create_docx():

    r = requests.post(instance_url+"/services/apexrest/data_retrieve/"+record_id,data=obj_wrapper,headers = {"Authorization":'OAuth '+session_id,'Content-Type': 'application/json','recordId':record_id})
    print("response-->{}".format(r))
    data_dict = json.loads(r.json())


    # r =  get_data_sf()
    # data_dict = json.loads(r.text)
    bind_values_doc(data_dict,doc)
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_bytes = docx_stream.getvalue()
    encoded = base64.b64encode(docx_bytes)
    data = {
            "Name" : "demoAttachment.pdf",
            "Body": str(encoded)[2:-1],
            "parentId": record_id 
            }
    # salesforce_response = requests.post(instance_url+"/services/data/v47.0/sobjects/Attachment",data=data,headers = {"Authorization":'OAuth '+'00D0p0000000V5HEAU!AQgAQEWkzrkC5T1NKhmV2C43BfSYR.2NbZ2VxatnomppLnH9V6hnFl1SCwIun9Cm2FI9Xbdpt_Lp6ie.pUyXXuFKfFisMEiJ','Content-Type': 'application/json'})
    # print("salesforce_response-->{}".format(salesforce_response))
    return json.dumps(data)

#Mock response, instead of salesforce response
@app.route('/bind_document', methods=['POST'])
def get_data_sf() :
    r = request.data
    data_dict = json.loads(r.text)
    print("request.data-->{}".format(data_dict))
    return requests.get("http://www.mocky.io/v2/5e4631fc3300004d00025f7f")

#To bind values to the fields which are not inside the table in the document
def bind_values_doc(data_dict,doc):
    # doc = docx.Document(file_path)
    # docume = Document(file_path)
    for paragraph in doc.paragraphs:
        matched_patterns = re.findall("\\$\\{(.*?)\\}", paragraph.text)
        function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", paragraph.text)
        if len(matched_patterns) > 0 :
            for value in matched_patterns :
                            text_in_cell = paragraph.text
                            field_value = attach_field_values(value,data_dict)
                            field_value = text_in_cell.replace('${'+value+'}',field_value)
                            paragraph.text = field_value
                            target_stream = StringIO()
        elif len(function_list) > 0 :
            field_value = ''
            field_value = generate_functions(function_list,data_dict)
            paragraph.text = field_value
        target_stream = StringIO()
        # doc.save(doc)
    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                    function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                    field_value = ''
                    if len(function_list) > 0 :
                            field_value = generate_functions(function_list,data_dict)
                            cell.text = field_value
                    elif len(matched_patterns) > 0 :
                        for value in matched_patterns :
                            text_in_cell = cell.text
                            field_value = attach_field_values(value,data_dict)
                            field_value = text_in_cell.replace('${'+value+'}',field_value)
                            cell.text = field_value
    target_stream = StringIO()
    
    # Iterating tables to bind parent field values
    if len(doc.tables) > 0:
        table_fields_list = []
        alltext_in_doc =[]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    alltext_in_doc.append(cell.text)
                    # cell.text = attach_field_values(cell.text,data_dict,file_path)
                    matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                    function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                    if len(matched_patterns) > 0 :
                        for value in matched_patterns :
                                cell.text = attach_field_values(value,data_dict)
                    elif len(function_list) > 0 :
                        field_value = ''
                        field_value = generate_functions(function_list,data_dict)
                        cell.text = field_value
                    target_stream = StringIO()
                    
                    # doc.save(doc)
        alltext_in_doc = '\n'.join(alltext_in_doc) 
        table_values = re.findall("\\$tbl\\{START:.*?\\}(.*?)\\$tbl\\{END:[A-Za-z]*\\}", alltext_in_doc.replace('\n', ' ').replace('\r', ''))
        if len(table_values) > 0 :
            table_fields_list = re.findall("\\$\\{(.*?)\\}",table_values[0])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_row_details = re.findall("\\$\\{(.*?)\\}", cell.text)
                        count_func_list = re.findall("\\{\\{RowCount:(.*?)\\}}", cell.text)
                        has_sum_func = re.findall("\\SUM\\{(.*?)\\}", cell.text)
                        if len(count_func_list) > 0 :
                            for value in count_func_list :
                                    text_in_cell = cell.text
                                    field_value = str(len(data_dict[count_func_list[0]]['records']))
                                    field_value = text_in_cell.replace('{{RowCount:'+value+'}}',field_value)
                                    cell.text = field_value
                        elif len(has_sum_func) > 0 :
                            splited_list = has_sum_func[0].split('.')
                            sum_of_field = 0 
                            for field in data_dict[splited_list[1]]['records'] : 
                                sum_of_field = sum_of_field + float(field[splited_list[2]])
                            text_in_cell = cell.text
                            field_value = str(sum_of_field)
                            field_value = text_in_cell.replace('$SUM{'+has_sum_func[0]+'}',field_value)
                            cell.text = field_value
                        
                        
                        if len(table_row_details) > 0 and table_row_details[0] not in table_fields_list :
                            matched_patterns = re.findall("\\$\\{(.*?)\\}", cell.text)
                            function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", cell.text)
                            field_value = ''
                            if len(function_list) > 0 :
                                    field_value = generate_functions(function_list,data_dict)
                                    cell.text = field_value
                            elif len(matched_patterns) > 0 :
                                for value in matched_patterns :
                                    text_in_cell = cell.text
                                    field_value = attach_field_values(value,data_dict)
                                    field_value = text_in_cell.replace('${'+value+'}',field_value)
                                    cell.text = field_value
            target_stream = StringIO()
            # r = requests.post("https://yourInstance.salesforce.com/services/data/v23.0/sobjects/ContentVersion",data=obj_wrapper,headers = 
            # curl https://yourInstance.salesforce.com/services/data/v23.0/sobjects/ContentVersion -H "Authorization: Bearer token" -H "Content-Type: multipart/form-data; boundary=\"boundary_string\"" --data-binary @NewContentVersion.json
            # doc.save(doc)
 
        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)          
        
        # Iterating tables to bind child field values       
        for table in doc.tables:
            column_value_list = []
            head_obj = []
            row_to_add = []
            for row_index,row in enumerate(table.rows) : 
                    for column_index,cell in enumerate(row.cells):
                            check_child =  re.findall("\\$\\{(.*?)\\}", cell.text)
                            if len(check_child) > 0 and  check_child[0] in  table_fields_list :
                                table_row_details = re.findall("\\$\\{(.*?)\\}", cell.text)
                            table_obj = re.findall("\\$tbl{START:(.*):", cell.text)
                            if len(table_obj) == 0:
                                table_obj = re.findall("\\$tbl\\{START:(.*?)\\}", cell.text)
                            table_end = re.findall("\\$tbl\\{END:(.*?)\\}", cell.text)
                            if len(table_obj) > 0:
                                row_to_add = table.row_cells(row_index)
                                head_obj = re.findall("\\$tbl{START:(.*):", cell.text)
                                if len(head_obj) == 0:
                                    head_obj = re.findall("\\$tbl\\{START:(.*?)\\}", cell.text)
                                head_obj[0] = head_obj[0].strip()
                                row_columns =[]
                                for cell in row_to_add:
                                    if cell.text not in row_columns :
                                        row_columns.append(cell.text)
                                if len(head_obj) > 0 :  
                                    if len(check_child) > 0 and  check_child[0] in  table_fields_list : 
                                        for i,record in enumerate(data_dict[head_obj[0]]['records']) : 
                                            current_row = table.rows[row_index]
                                            border_copied = copy.deepcopy(current_row._tr)
                                            tr = border_copied
                                            current_row._tr.addnext(tr)
                                            for j,column in enumerate(row_columns):
                                                table_pattern  = re.findall("\\$\\{(.*?)\\}", column)
                                                if len(table_pattern) > 0 :
                                                    field_pattern = table_pattern[0].split('.')
                                                    table.cell(row_index+1, j).text = record[field_pattern[-1]]
                                                    # table.rows[row_index+1].height_rule = WD_ROW_HEIGHT.AUTO
                                                    # print("table-->{}".format(table.rows[row_index+1].height_rule))
                                                    table.rows[row_index+1].height = 1        
                            if len(table_end) > 0 :
                                remove_row(table, table.rows[row_index])
        # doc.save(doc) 
    
#Method to get field index
#Parameters (fieldName, metaData, objName)
def get_field_index(field_name, data,list_name,obj_name):
        obj_list = list()
        for record in data[list_name]:
            obj_list.append(record[obj_name])
        if len(obj_list) > 0:
            try:
                return obj_list.index(field_name)
            except ValueError:
                return -1
        else: 
            return -1     

#Method to manipulate functions in the document
def generate_functions(function_list,data_dict) :
    if_condition_list = re.findall("IF\\((.*?)\\)", function_list[0])
    if len(if_condition_list) > 0 :
        conditon_value,true_value,false_value = if_condition_list[0].split(',')[0],if_condition_list[0].split(',')[1],if_condition_list[0].split(',')[2]
        field_name_list = re.findall("\\$\\{(.*?)\\}", conditon_value)
        conv_value_to_str = re.split(' ',conditon_value)
        added_changes = conditon_value.replace(conv_value_to_str[-1],"'"+conv_value_to_str[-1]+"'")
        if len(field_name_list) > 0 :
                splited_list = field_name_list[0].split('.')
                if len(splited_list) == 2 :
                    field_value = data_dict[splited_list[1]]
                elif len(splited_list) == 3 :
                    obj_name_match = re.split('__r',splited_list[1])
                    field_value = data_dict[obj_name_match[0]][splited_list[2]]
                elif len(splited_list) == 4 :
                    parent_name_match = re.split('__r',splited_list[1])
                    grand_name_match = re.split('__r',splited_list[2])
                    field_value = data_dict[parent_name_match[0]][grand_name_match[0]][splited_list[3]]
        field_value = field_value.replace(" ","")
        val = added_changes.replace('${'+field_name_list[0]+'}',"'"+field_value.strip()+"'")
        cons = eval("true_value if "+val+" else false_value")
        return cons
    else : 
        return "Error"


#To bind values from salesforce to the matched string
#Parameters(fieldName, metaData, filePath)
def attach_field_values(obj_to_bind,data_dict) :
     function_list = re.findall("\\{\\{FUNC:(.*?)\\}}", obj_to_bind)
     field_name = ''
     if len(function_list) > 0 :
            field_name = generate_functions(function_list,data_dict)
     else :
            format_type = re.findall("(#[A-Z]*)",obj_to_bind)
            corrected_field = ''
            if len(format_type) > 0 :
                corrected_field = obj_to_bind.replace(format_type[0],'').rstrip()
            else :
                corrected_field = obj_to_bind
            splited_list = corrected_field.split('.')

            if len(splited_list) == 2 :
                print("splited_list-->{}".format(splited_list[1]))
                print('data_dict_type--> {}'.format(type(data_dict)))
                formatted_type = data_dict[splited_list[1]]
                if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                    value = '$'+value
                elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    value = date(int(separate_date[2]), int(separate_date[1]), int(separate_date[0])).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                field_name = value if len(format_type) > 0 else formatted_type
            elif len(splited_list) == 3 :
                obj_name_match = re.split('__r',splited_list[1])
                formatted_type = data_dict[obj_name_match[0]][splited_list[2]]
                if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                    value = '$'+value
                elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    value = date(int(separate_date[2]), int(separate_date[1]), int(separate_date[0])).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                field_name = value if len(format_type) > 0 else formatted_type
            elif len(splited_list) == 4 :
                 parent_name_match = re.split('__r',splited_list[1])
                 grand_name_match = re.split('__r',splited_list[2])
                 formatted_type = data_dict[parent_name_match[0]][grand_name_match[0]][splited_list[3]]
                 if len(format_type) > 0 and format_type[0] == '#NUMBER' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                 elif len(format_type) > 0 and format_type[0] == '#CURRENCY' :
                    value = ','.join(formatted_type[i:i+3] for i in range(0, len(formatted_type), 3))
                    value = '$'+value
                 elif len(format_type) > 0 and format_type[0] == '#DATE' :
                    separate_date = formatted_type.split('-')
                    value = date(int(separate_date[2]), int(separate_date[1]), int(separate_date[0])).ctime()
                    value = value.split(' ')
                    value = value[1]+' '+value[2]+','+''+value[-1]
                 field_name = value if len(format_type) > 0 else formatted_type

     return field_name

if __name__ == "__main__":
    app.run(host='localhost', port=8080, debug=True)
    


